import os
import re
from docx import Document
from docx.shared import Pt

ARTICLES_DIR = "content/articles"
OUTPUT_DIR = "/Users/desktop/Desktop"

def convert_mdx_to_docx():
    if not os.path.exists(ARTICLES_DIR):
        print(f"Directory {ARTICLES_DIR} does not exist.")
        return

    files = [f for f in os.listdir(ARTICLES_DIR) if f.endswith('.mdx')]
    
    for filename in files:
        filepath = os.path.join(ARTICLES_DIR, filename)
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract title from frontmatter
        title_match = re.search(r"title:\s*['\"]?(.*?)['\"]?\n", content)
        title = title_match.group(1) if title_match else filename.replace('.mdx', '')

        # Strip frontmatter
        content = re.sub(r"^---\n.*?\n---\n", "", content, flags=re.DOTALL)
        
        # Clean up some mdx components if any (like <Callout> etc)
        content = re.sub(r"<[^>]+>", "", content)

        doc = Document()
        
        # Add title
        heading = doc.add_heading(title, 0)
        
        # Add paragraphs
        for paragraph in content.split('\n\n'):
            p_text = paragraph.strip()
            if not p_text:
                continue
            
            # Simple bold removal/conversion (python-docx requires adding runs for inline bold)
            # For simplicity, we just strip the ** but keep the text
            p_text = p_text.replace('**', '')
            p_text = p_text.replace('__', '')
            p_text = p_text.replace('`', '')
            
            if p_text.startswith('#'):
                # Handle headings
                level_match = re.match(r"^(#+)\s*(.*)", p_text)
                if level_match:
                    level = len(level_match.group(1))
                    text = level_match.group(2)
                    doc.add_heading(text, level=min(level, 9))
            else:
                doc.add_paragraph(p_text)
        
        out_name = filename.replace('.mdx', '.docx')
        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)
        print(f"Generated: {out_path}")

if __name__ == "__main__":
    convert_mdx_to_docx()
